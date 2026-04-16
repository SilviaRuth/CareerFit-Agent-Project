"""Bounded file ingestion for txt, pdf, and docx inputs."""

from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.core.config import SUPPORTED_INGESTION_EXTENSIONS
from app.schemas.parse import ParserDiagnostic


@dataclass(slots=True)
class IngestedDocument:
    """Raw text plus source metadata captured during ingestion."""

    raw_text: str
    source_type: str
    source_name: str | None = None
    media_type: str | None = None
    warnings: list[ParserDiagnostic] = field(default_factory=list)


def ingest_text(text: str, source_name: str | None = None) -> IngestedDocument:
    """Wrap direct text input with consistent ingestion metadata."""
    warnings: list[ParserDiagnostic] = []
    if not text.strip():
        warnings.append(
            ParserDiagnostic(
                warning_code="empty_text_input",
                message="Input text is empty; extraction will be partial.",
                section=None,
                severity="error",
                source="ingestion",
            )
        )

    return IngestedDocument(
        raw_text=text,
        source_type="text",
        source_name=source_name,
        media_type="text/plain",
        warnings=warnings,
    )


def ingest_file(content: bytes, filename: str, media_type: str | None = None) -> IngestedDocument:
    """Read supported file bytes into raw text with bounded diagnostics."""
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_INGESTION_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_INGESTION_EXTENSIONS))
        raise ValueError(f"Unsupported file type '{extension or '<none>'}'. Expected one of {supported}.")

    warnings: list[ParserDiagnostic] = []
    if extension == ".txt":
        raw_text, decode_warnings = _read_txt(content)
        warnings.extend(decode_warnings)
    elif extension == ".pdf":
        raw_text, decode_warnings = _read_pdf(content)
        warnings.extend(decode_warnings)
    else:
        raw_text, decode_warnings = _read_docx(content)
        warnings.extend(decode_warnings)

    if not raw_text.strip():
        warnings.append(
            ParserDiagnostic(
                warning_code="empty_extracted_text",
                message="The file produced little or no text during ingestion.",
                section=None,
                severity="error",
                source="ingestion",
            )
        )

    return IngestedDocument(
        raw_text=raw_text,
        source_type="file",
        source_name=filename,
        media_type=media_type or SUPPORTED_INGESTION_EXTENSIONS[extension],
        warnings=warnings,
    )


def _read_txt(content: bytes) -> tuple[str, list[ParserDiagnostic]]:
    """Decode a text file with bounded fallback behavior."""
    try:
        return content.decode("utf-8-sig"), []
    except UnicodeDecodeError:
        return (
            content.decode("latin-1"),
            [
                ParserDiagnostic(
                    warning_code="txt_decode_fallback",
                    message="TXT input required latin-1 fallback decoding.",
                    section=None,
                    severity="warning",
                    source="ingestion",
                )
            ],
        )


def _read_pdf(content: bytes) -> tuple[str, list[ParserDiagnostic]]:
    """Extract text from a PDF without OCR or layout reconstruction."""
    warnings: list[ParserDiagnostic] = []
    reader = PdfReader(BytesIO(content))
    lines: list[str] = []

    for page_index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            lines.append(page_text.strip())
            continue
        warnings.append(
            ParserDiagnostic(
                warning_code="pdf_page_without_text",
                message=f"PDF page {page_index} did not yield extractable text.",
                section=None,
                severity="warning",
                source="ingestion",
            )
        )

    if reader.is_encrypted:
        warnings.append(
            ParserDiagnostic(
                warning_code="pdf_encrypted",
                message="PDF appears encrypted; extracted text may be incomplete.",
                section=None,
                severity="warning",
                source="ingestion",
            )
        )

    return "\n\n".join(lines), warnings


def _read_docx(content: bytes) -> tuple[str, list[ParserDiagnostic]]:
    """Extract plain paragraph text from a DOCX document."""
    warnings: list[ParserDiagnostic] = []
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    table_cell_count = sum(len(row.cells) for table in document.tables for row in table.rows)
    if table_cell_count:
        warnings.append(
            ParserDiagnostic(
                warning_code="docx_tables_ignored",
                message="DOCX tables were ignored by the bounded parser.",
                section=None,
                severity="info",
                source="ingestion",
            )
        )

    return "\n\n".join(paragraphs), warnings
